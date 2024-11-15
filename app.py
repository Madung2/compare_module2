import gradio as gr
from docx import Document
import json
from lxml import etree
import zipfile
import pandas as pd
from transformers import AutoTokenizer, AutoModelForTokenClassification, pipeline
import re
from bs4 import BeautifulSoup
# Initialize NER pipeline
model_name = "xlm-roberta-large-finetuned-conll03-english"
tokenizer = AutoTokenizer.from_pretrained(f"FacebookAI/{model_name}")
model = AutoModelForTokenClassification.from_pretrained(f"FacebookAI/{model_name}")
ner_pipeline = pipeline("ner", model=model, tokenizer=tokenizer, aggregation_strategy="simple")

def highlight_target_in_html(html_content, target):
    # Step 1: 기본적으로 모든 `target`을 하이라이트
    highlighted_html = html_content.replace(target, f"<mark style='background-color: yellow'>{target}</mark>")
    
    # Step 2: BeautifulSoup을 사용하여 행 단위로 처리
    soup = BeautifulSoup(highlighted_html, 'html.parser')
    rows = soup.find_all('tr')
    for row in rows:
        cells = row.find_all('td')  # 각 행의 모든 셀을 찾음
        # 첫 번째 셀에 하이라이트가 있는지 검사
        if cells and '<mark style=' in cells[0].decode_contents():
            # 모든 셀에 하이라이트 적용
            for cell in cells:
                cell_contents = cell.decode_contents()
                # 기존에 이미 하이라이트가 적용되지 않은 경우 추가로 하이라이트 적용
                if '<mark style=' not in cell_contents:
                    cell.string = cell_contents.replace(target, f"<mark style='background-color: yellow'>{target}</mark>")
    return str(soup)
def highlight_text(text):
    return f"<mark style='background-color: yellow'>{text}</mark>"

def highlight_ner(text, res_type):
    print('highlight:', text)
    ner_results = ner_pipeline(text)
    highlighted_text = text
    for entity in ner_results:
        if entity['entity_group'] == res_type:
            highlighted_text = highlighted_text.replace(entity['word'], highlight_text(entity['word']))
    return highlighted_text

def run_ner(text, res_type):
    # 지금은 첫번째것만 리턴하고 있음
    ner_results = ner_pipeline(text)
    for entity in ner_results:
        if entity['entity_group'] == res_type:
            return entity['word']

def reorganize_result(result):
    reorganized = []
    for item in result:
        if item['type'] == 'table' and reorganized:
            last_item = reorganized[-1]
            if last_item['type'] == 'text':
                last_item['table'] = item['res']
            else:
                reorganized.append(item)
        else:
            reorganized.append(item)
    return reorganized

def create_text_table_mapping(result):
    mapping = {}
    previous_text = None
    for item in result:
        if item['type'] == 'text':
            previous_text = item['res']
        elif item['type'] == 'table' and previous_text is not None:
            mapping[previous_text] = item['res']
            previous_text = None
    return mapping
def extract_table_mapping(doc_order):
    """
    This function takes a list of dictionaries representing document structures,
    and extracts key-value pairs where keys are the 'out' values of the 'para'
    type entries right before 'table' type entries, and values are the 'out' lists of the 'table'.
    """
    table_mapping = {}
    for idx, entry in enumerate(doc_order):
        if entry['type'] == 'table' and idx > 0:
            # Find the preceding 'para' key as the key for this table
            prev_entry = doc_order[idx - 1]
            if prev_entry['type'] == 'para' and prev_entry['out']:
                key = prev_entry['out']
                value = entry['out']
                table_mapping[key] = value
    return table_mapping
def read_opinion(file_path):
    doc = Document(file_path)
    # Access the underlying XML tree
    xml_content = doc.element.body

    # Define namespaces
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    table_idx =0
    # Iterate over elements to find paragraphs and tables in order
    doc_order = []
    for i, element in enumerate(xml_content):
        if element.tag.endswith('p'):
            # This is a paragraph
            para = element
            # Extract paragraph text
            texts = para.findall('.//w:t', namespaces)
            paragraph_text = ''.join([text.text for text in texts if text is not None])
            print(f"Element {i} - Paragraph: {paragraph_text}")
            doc_order.append({'type':'para', 'out':paragraph_text})
        elif element.tag.endswith('tbl'):
            print(f"Element {i} - Table: {table_idx}")
            table = doc.tables[table_idx]
            table_dict=[]
            for row in table.rows:
                row_t = [cell.text for cell in row.cells]
                table_dict.append(row_t)
                print(row_t)
            doc_order.append({'type':'table', 'out': table_dict, 'table_idx':table_idx})
            # doc.tables[table_idx]._element.getparent().remove(doc.tables[table_idx]._element)
            table_idx+=1
    res = extract_table_mapping(doc_order)
    # with zipfile.ZipFile(file_path) as docx_zip:
    #     with docx_zip.open('word/document.xml') as xml_file:
    #         tree = etree.parse(xml_file)
    #         root = tree.getroot()
    #         nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    #         result = []
    #         in_table = False
    #         for elem in root.iter():
    #             if elem.tag == f"{{{nsmap['w']}}}tbl":
    #                 in_table = True
    #                 table_data = []
    #                 for row in elem.findall('.//w:tr', namespaces=nsmap):
    #                     row_data = [cell.text for cell in row.findall('.//w:t', namespaces=nsmap) if cell.text]
    #                     table_data.append(row_data)
    #                 result.append({'type': 'table', 'res': table_data})
    #                 in_table = False
    #             elif elem.tag == f"{{{nsmap['w']}}}p" and not in_table:
    #                 text = ''.join(e.text for e in elem.iter() if e.text).strip()
    #                 if text:
    #                     result.append({'type': 'text', 'res': text})
    # res = create_text_table_mapping(result)
    return res

def read_docx(file):
    doc = Document(file)
    content = {}
    current_title = None
    for para in doc.paragraphs:
        if para.text.strip():
            is_all_bold = all(run.bold for run in para.runs if run.text.strip())
            if is_all_bold:
                current_title = para.text.strip()
                content[current_title] = []
            elif current_title:
                content[current_title].append(para.text.strip())
    content = {k: [v for v in values if v] for k, values in content.items() if k}
    return content

def compare_texts(text1, text2, json_input):
    res = []
    for json_key, v in json_input.items():
        contract_block = v['contract']['block']
        opinion_block = v['opinion']['block']
        contract_syn = v['contract']['syn']
        opinion_syn = v['opinion']['syn']
        res_type = v['contract'].get('res', None)  # Get the res type if present

        con_res = {}
        for k, v in text1.items():
            if any(b in k for b in contract_block):
                for line in v:
                    if any(s in line for s in contract_syn):
                        if res_type:
                            ner_res = run_ner(line, res_type) # 지금은 res값이 1개
                            con_res = {'all_text': v, 'target': line, 'sp_target':ner_res, 'title':k}
                        else:
                            con_res = {'all_text': v, 'target': line, 'title':k}
                        break
        op_res = {}
        for key, table in text2.items():

            if any(b in key for b in opinion_block):
                df = pd.DataFrame(table)
                html_table = df.to_html(index=False, escape=False)
                for s in opinion_syn:
                    if s in html_table:
                        op_res = {'all_text': html_table, 'target': s, 'title':key}
        res.append([con_res, op_res, json_key])
    return res

def highlight_target_in_html(html_content, target):
    highlighted_html = html_content.replace(
        target, f"<mark style='background-color: yellow'>{target}</mark>")
    return highlighted_html
def highlight_target_in_html(html_content, target):
    # Step 1: 기본적으로 모든 `target`을 하이라이트
    highlighted_html = html_content.replace(target, f"<mark style='background-color: yellow'>{target}</mark>")
    
    # Step 2: BeautifulSoup을 사용하여 행 단위로 처리
    soup = BeautifulSoup(highlighted_html, 'html.parser')
    rows = soup.find_all('tr')
    for row in rows:
        cells = row.find_all('td')  # 각 행의 모든 셀을 찾음
        # 첫 번째 셀에 하이라이트가 있는지 검사
        if cells and '<mark style=' in cells[0].decode_contents():
            # 모든 셀에 하이라이트 적용
            for cell in cells:
                cell_contents = cell.decode_contents()
                # 셀의 기존 내용을 전체 하이라이트 처리
                cell.clear()
                cell.append(BeautifulSoup(f"<mark style='background-color: yellow'>{cell_contents}</mark>", 'html.parser'))
    return str(soup)
def process_files(agreement_file, opinion_file, json_input):
    try:
        json_input = json.loads(json_input)
    except json.JSONDecodeError:
        return "올바른 JSON 형식이 아닙니다. 입력을 확인하세요.", ""

    agreement_text = read_docx(agreement_file)
    opinion_text = read_opinion(opinion_file)
    comparison_result = compare_texts(agreement_text, opinion_text, json_input)
    agreement_output = ""
    opinion_output = ""

    for index, (con, op, rep) in enumerate(comparison_result):
        agreement_output += f"\n ### 약정서 내용 (섹션 {index + 1}: {con['title']})\n"
        if isinstance(con, dict) and 'all_text' in con and 'target' in con:
            for idx, text in enumerate(con['all_text'], start=1):
                if text == con['target']:
                    if 'sp_target' in con:
                        # ner 값만 하이라이트
                        highlighted_target = highlight_target_in_html(text, con['sp_target'])
                        agreement_output += f"{idx}. {highlighted_target}\n\n"
                    else:
                        highlighted_target = highlight_text(text)
                        agreement_output += f"{idx}. {highlighted_target}\n\n"
                else:
                    agreement_output += f"{idx}. {text}\n\n"
        else:
            agreement_output += f"{con}\n\n"

        opinion_output += f"\n\n### 의견서 내용 (섹션 {index + 1}: {op['title']})\n"
        if isinstance(op, dict) and 'all_text' in op:
            if op['target'] in op['all_text']:
                highlighted_html = highlight_target_in_html(op['all_text'], op['target'])
                opinion_output += highlighted_html            
            else:
                opinion_output += op['all_text']


    # 맞춤을 위해 zip_longest 사용하여 길이 조정

################################################################

    return agreement_output, opinion_output

# Gradio interface
json_input = """
{
    "차주" : {
        "contract": {"block":["대출계약서"], "syn":["차주"], "res":"ORG"},
        "opinion" : {"block":["신청내용"], "syn":["차주"], "res":"ORG"}
    },
    "상환방법" : {
        "contract": {"block":["대출금의 상환"], "syn":["일시", "분할"] },
        "opinion" : {"block":["신청내용"], "syn":["만기"] }
    },
    "대리금융기관" : {
        "contract": {"block":["대출계약서"], "syn":["대리금융기관"], "res":"ORG" },
        "opinion" : {"block":["신청내용"], "syn":["대리금융기관"], "res":"ORG" }
    },
    "대출금액" : {
        "contract": {"block":["대출계약서"], "syn":["조달액"]},
        "opinion" : {"block":["신청내용"], "syn":["조달금액"]}
    }
}
"""

with gr.Blocks() as demo:
    gr.Markdown("## 약정서-의견서 비교")
    
    with gr.Row():
        agreement_file = gr.File(label="약정서 파일 업로드 (docx)")
        opinion_file = gr.File(label="의견서 파일 업로드 (docx)")
    
    json_input_box = gr.Textbox(label="JSON 비교 기준 입력", value=json_input, lines=10)
    
    compare_button = gr.Button("비교 실행")
    
    with gr.Row():
        # final_output = gr.Markdown(label= "결과")
        agreement_output = gr.Markdown(label="약정서 내용")
        opinion_output = gr.Markdown(label="의견서 내용")
    
    compare_button.click(
        process_files,
        inputs=[agreement_file, opinion_file, json_input_box],
        # outputs = [final_output]
        outputs=[agreement_output, opinion_output]
    )
demo.launch(server_name="0.0.0.0", server_port=7860)
